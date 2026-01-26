"""Document parsing utilities for extracting text and diagrams from various file formats."""
import io
import os
from typing import Optional, Tuple
from fastapi import UploadFile, HTTPException
from pypdf import PdfReader
from docx import Document
import openpyxl

from diagram_extraction import (
    DiagramExtractionResult,
    extract_diagrams_with_document_intelligence,
    extract_diagrams_from_path as _extract_diagrams_from_path,
)


MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB (increased for async processing)

# Supported image extensions for Azure AI Document Intelligence
IMAGE_EXTENSIONS = {"png", "jpg", "jpeg", "gif", "bmp", "tiff", "tif", "webp"}

# File types that support diagram extraction (PDFs and images)
DIAGRAM_EXTRACTABLE_EXTENSIONS = {"pdf"} | IMAGE_EXTENSIONS


def _is_azure_document_intelligence_configured() -> bool:
    """Check if Azure AI Document Intelligence is configured."""
    endpoint = os.getenv("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT")
    key = os.getenv("AZURE_DOCUMENT_INTELLIGENCE_KEY")
    return bool(endpoint and key)


def _extract_with_document_intelligence(content: bytes, filename: str = "") -> str:
    """
    Extract text from document using Azure AI Document Intelligence.

    Uses the prebuilt-read model for OCR (Optical Character Recognition).
    Works with images, scanned PDFs, and PDFs with embedded images.

    Args:
        content: Document file bytes (image or PDF)
        filename: Original filename for logging

    Returns:
        Extracted text from the document

    Raises:
        ValueError: If Azure AI Document Intelligence is not configured or extraction fails
    """
    import base64

    endpoint = os.getenv("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT")
    key = os.getenv("AZURE_DOCUMENT_INTELLIGENCE_KEY")

    if not endpoint or not key:
        raise ValueError(
            "Azure AI Document Intelligence is not configured. "
            "Set AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT and AZURE_DOCUMENT_INTELLIGENCE_KEY environment variables."
        )

    try:
        from azure.ai.documentintelligence import DocumentIntelligenceClient
        from azure.ai.documentintelligence.models import AnalyzeDocumentRequest
        from azure.core.credentials import AzureKeyCredential

        # Create client
        client = DocumentIntelligenceClient(
            endpoint=endpoint,
            credential=AzureKeyCredential(key)
        )

        # Base64 encode the content as required by the API
        # See: https://learn.microsoft.com/en-us/python/api/azure-ai-documentintelligence/azure.ai.documentintelligence.models.analyzedocumentrequest
        base64_content = base64.b64encode(content)

        # Analyze document using prebuilt-read model (best for OCR)
        poller = client.begin_analyze_document(
            model_id="prebuilt-read",
            analyze_request=AnalyzeDocumentRequest(bytes_source=base64_content),
        )

        result = poller.result()

        # Extract text from all pages
        text_parts = []
        if result.content:
            text_parts.append(result.content)

        extracted_text = "\n".join(text_parts)

        if not extracted_text.strip():
            raise ValueError("No text could be extracted from the document")

        print(f"  📄 Azure AI extracted {len(extracted_text)} characters from {filename}")
        return extracted_text

    except ImportError:
        raise ValueError(
            "azure-ai-documentintelligence package is not installed. "
            "Run: pip install azure-ai-documentintelligence"
        )
    except Exception as e:
        raise ValueError(f"Failed to extract text with Azure AI Document Intelligence: {str(e)}")


async def extract_text_from_file(file: UploadFile) -> str:
    """
    Extract text content from uploaded file.

    Supports: PDF, DOCX, TXT, MD, XLSX, images (PNG, JPG, JPEG, GIF, BMP, TIFF, WEBP).

    When Azure AI Document Intelligence is configured:
    - Images: OCR extracts text from images
    - PDFs: Enhanced extraction handles scanned PDFs and embedded images
    - Falls back to pypdf for PDFs if Azure AI fails

    When Azure AI Document Intelligence is NOT configured:
    - Images: Will return an error (requires Azure AI)
    - PDFs: Uses pypdf (text-based PDFs only, no OCR for scanned pages)

    Args:
        file: Uploaded file from FastAPI

    Returns:
        Extracted text content

    Raises:
        HTTPException: If file is too large, empty, or cannot be processed
    """
    # Read file content
    content = await file.read()
    file_size = len(content)

    # Validate file size
    if file_size > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413,
            detail=f"File too large. Maximum size is {MAX_FILE_SIZE / 1024 / 1024}MB"
        )

    if file_size == 0:
        raise HTTPException(status_code=400, detail="File is empty")

    # Get file extension
    filename = file.filename or ""
    extension = filename.lower().split(".")[-1] if "." in filename else ""

    try:
        # Extract text based on file type
        if extension == "pdf":
            text = _extract_from_pdf(content, filename)
        elif extension in ["docx", "doc"]:
            text = _extract_from_docx(content)
        elif extension in ["xlsx", "xls"]:
            text = _extract_from_excel(content)
        elif extension in ["txt", "md", "markdown", "text"]:
            text = _extract_from_text(content)
        elif extension in IMAGE_EXTENSIONS:
            # Image files require Azure AI Document Intelligence
            if not _is_azure_document_intelligence_configured():
                raise HTTPException(
                    status_code=400,
                    detail="Image upload requires Azure AI Document Intelligence to be configured. "
                           "Please configure AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT and AZURE_DOCUMENT_INTELLIGENCE_KEY."
                )
            text = _extract_with_document_intelligence(content, filename)
        else:
            # Try to read as plain text
            text = _extract_from_text(content)

        # Validate extracted content
        if not text or not text.strip():
            raise HTTPException(
                status_code=400,
                detail="Could not extract any text from the document"
            )

        return text.strip()

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Error processing file: {str(e)}"
        )


def _extract_from_pdf(content: bytes, filename: str = "") -> str:
    """
    Extract text from PDF file.

    If Azure AI Document Intelligence is configured, uses it for better extraction
    of scanned PDFs and embedded images. Falls back to pypdf otherwise.

    Args:
        content: PDF file bytes
        filename: Original filename for logging

    Returns:
        Extracted text from the PDF
    """
    # Try Azure AI Document Intelligence first if configured
    # This handles scanned PDFs and embedded images better than pypdf
    if _is_azure_document_intelligence_configured():
        try:
            return _extract_with_document_intelligence(content, filename)
        except Exception as e:
            print(f"  ⚠️ Azure AI Document Intelligence failed, falling back to pypdf: {e}")
            # Fall through to pypdf extraction

    # Fall back to pypdf for text-based PDFs
    pdf_file = io.BytesIO(content)
    reader = PdfReader(pdf_file)

    text_parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)

    return "\n\n".join(text_parts)


def _extract_from_docx(content: bytes) -> str:
    """Extract text from DOCX file."""
    docx_file = io.BytesIO(content)
    doc = Document(docx_file)

    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text)

    return "\n\n".join(text_parts)


def _extract_from_excel(content: bytes) -> str:
    """Extract text from Excel file."""
    excel_file = io.BytesIO(content)
    workbook = openpyxl.load_workbook(excel_file, data_only=True)

    text_parts = []
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join(str(cell) for cell in row if cell is not None)
            if row_text.strip():
                text_parts.append(row_text)

    return "\n".join(text_parts)


def _extract_from_text(content: bytes) -> str:
    """Extract text from plain text file."""
    # Try different encodings
    encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]

    for encoding in encodings:
        try:
            return content.decode(encoding)
        except (UnicodeDecodeError, AttributeError):
            continue

    raise ValueError("Could not decode text file with any supported encoding")


def extract_text_from_path(file_path: str) -> str:
    """
    Extract text content from a file path (for async background processing).

    Supports: PDF, DOCX, TXT, MD, XLSX, images (PNG, JPG, JPEG, GIF, BMP, TIFF, WEBP).

    When Azure AI Document Intelligence is configured:
    - Images: OCR extracts text from images
    - PDFs: Enhanced extraction handles scanned PDFs and embedded images
    - Falls back to pypdf for PDFs if Azure AI fails

    When Azure AI Document Intelligence is NOT configured:
    - Images: Will return an error (requires Azure AI)
    - PDFs: Uses pypdf (text-based PDFs only, no OCR for scanned pages)

    Args:
        file_path: Absolute path to the file

    Returns:
        Extracted text content

    Raises:
        ValueError: If file cannot be processed or is empty
    """
    if not os.path.exists(file_path):
        raise ValueError(f"File not found: {file_path}")

    file_size = os.path.getsize(file_path)
    if file_size == 0:
        raise ValueError("File is empty")

    if file_size > MAX_FILE_SIZE:
        raise ValueError(f"File too large. Maximum size is {MAX_FILE_SIZE / 1024 / 1024}MB")

    # Get file extension and filename
    extension = file_path.lower().split(".")[-1] if "." in file_path else ""
    filename = os.path.basename(file_path)

    # Read file content
    with open(file_path, "rb") as f:
        content = f.read()

    try:
        # Extract text based on file type
        if extension == "pdf":
            text = _extract_from_pdf(content, filename)
        elif extension in ["docx", "doc"]:
            text = _extract_from_docx(content)
        elif extension in ["xlsx", "xls"]:
            text = _extract_from_excel(content)
        elif extension in ["txt", "md", "markdown", "text"]:
            text = _extract_from_text(content)
        elif extension in IMAGE_EXTENSIONS:
            # Image files require Azure AI Document Intelligence
            if not _is_azure_document_intelligence_configured():
                raise ValueError(
                    "Image upload requires Azure AI Document Intelligence to be configured. "
                    "Please configure AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT and AZURE_DOCUMENT_INTELLIGENCE_KEY."
                )
            text = _extract_with_document_intelligence(content, filename)
        else:
            # Try to read as plain text
            text = _extract_from_text(content)

        # Validate extracted content
        if not text or not text.strip():
            raise ValueError("Could not extract any text from the document")

        return text.strip()

    except Exception as e:
        raise ValueError(f"Error processing file: {str(e)}")


def extract_text_and_diagrams_from_path(
    file_path: str,
    analyze_with_vision: bool = True
) -> Tuple[str, DiagramExtractionResult]:
    """
    Extract both text and diagrams from a file path.

    For PDF and image files, uses Azure Document Intelligence with the
    prebuilt-layout model to detect figures and extract text. Optionally
    analyzes diagrams with GPT-4 Vision.

    For other file types, extracts text only (no diagram detection).

    Args:
        file_path: Absolute path to the file
        analyze_with_vision: Whether to analyze diagrams with GPT-4 Vision

    Returns:
        Tuple of (extracted_text, diagram_extraction_result)
    """
    if not os.path.exists(file_path):
        raise ValueError(f"File not found: {file_path}")

    file_size = os.path.getsize(file_path)
    if file_size == 0:
        raise ValueError("File is empty")

    if file_size > MAX_FILE_SIZE:
        raise ValueError(f"File too large. Maximum size is {MAX_FILE_SIZE / 1024 / 1024}MB")

    extension = file_path.lower().split(".")[-1] if "." in file_path else ""
    filename = os.path.basename(file_path)

    # For diagram-extractable files (PDF and images), use diagram extraction
    if extension in DIAGRAM_EXTRACTABLE_EXTENSIONS:
        if not _is_azure_document_intelligence_configured():
            # Fall back to text-only extraction for PDFs
            if extension == "pdf":
                text = extract_text_from_path(file_path)
                return text, DiagramExtractionResult(
                    document_name=filename,
                    extraction_method="text-only",
                    text_content=text,
                    diagram_summary="Azure AI Document Intelligence not configured. Diagram extraction requires configuration."
                )
            else:
                raise ValueError(
                    "Image and diagram extraction requires Azure AI Document Intelligence. "
                    "Please configure AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT and AZURE_DOCUMENT_INTELLIGENCE_KEY."
                )

        # Use diagram extraction (which also extracts text)
        result = _extract_diagrams_from_path(file_path, analyze_with_vision)

        # Check if diagram extraction failed - fall back to text-only extraction
        if result.extraction_method == "error" or not result.text_content.strip():
            print(f"  ⚠️ Diagram extraction failed, falling back to text-only extraction")
            try:
                # For PDFs, fall back to pypdf extraction
                if extension == "pdf":
                    text = extract_text_from_path(file_path)
                    return text, DiagramExtractionResult(
                        document_name=filename,
                        extraction_method="text-only-fallback",
                        text_content=text,
                        diagram_summary=f"Diagram extraction failed ({result.diagram_summary}). Used fallback text extraction."
                    )
                else:
                    # For images, we can't fall back - raise an error
                    raise ValueError(
                        f"Failed to extract content from image: {result.diagram_summary}"
                    )
            except Exception as fallback_error:
                raise ValueError(
                    f"Diagram extraction failed and fallback also failed: {result.diagram_summary}. "
                    f"Fallback error: {str(fallback_error)}"
                )

        return result.text_content, result

    # For other file types, extract text only
    text = extract_text_from_path(file_path)
    return text, DiagramExtractionResult(
        document_name=filename,
        extraction_method="text-only",
        text_content=text,
        diagram_summary="File type does not support diagram extraction."
    )


def supports_diagram_extraction(filename: str) -> bool:
    """
    Check if a file type supports diagram extraction.

    Args:
        filename: Filename to check

    Returns:
        True if the file type supports diagram extraction
    """
    extension = filename.lower().split(".")[-1] if "." in filename else ""
    return extension in DIAGRAM_EXTRACTABLE_EXTENSIONS


def get_diagram_extraction_status() -> dict:
    """
    Get the current status of diagram extraction capabilities.

    Returns:
        Dictionary with configuration status
    """
    from diagram_extraction import _is_vision_analysis_configured

    return {
        "document_intelligence_configured": _is_azure_document_intelligence_configured(),
        "vision_analysis_configured": _is_vision_analysis_configured(),
        "supported_extensions": list(DIAGRAM_EXTRACTABLE_EXTENSIONS),
        "capabilities": {
            "text_extraction": True,
            "diagram_detection": _is_azure_document_intelligence_configured(),
            "diagram_analysis": _is_vision_analysis_configured(),
        }
    }
